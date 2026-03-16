#!/usr/bin/env python3
"""
LDA Auto-Pipeline Runner

Runs every 5 minutes via LaunchAgent. For each completed LDA batch:
1. If batch has saved instructions -> run CC pipeline -> deliver results via Signal
2. If no instructions -> send notification to user

Pipeline: group by folder -> CC (all docs in folder together) -> deanonymize -> format -> deliver
Supports related_batches for cross-batch grouping into a single unified memo.

Task types:
- memo (default): Legal memo drafting using Office Memo template
- review-comments: Contract review with inline comments as blockquotes
- review-redline: Contract review with tracked changes (strikethrough/bold markup)
- draft: Contract drafting from templates (LOCAL Ollama only, no cloud API)
"""

import fcntl
import json
import os
import re
import subprocess
import sys
import time
import urllib.request
from collections import defaultdict
from datetime import datetime
from pathlib import Path

JOBS_DIR = Path.home() / ".openclaw/tools/lda/jobs"
META_DIR = JOBS_DIR / ".batch_meta"
OUTPUT_DIR = Path.home() / ".openclaw/tools/lda/output"
LOCK_FILE = JOBS_DIR / ".pipeline.lock"
PROCESSED_FILE = JOBS_DIR / ".processed_batches"
NOTIFIED_FILE = JOBS_DIR / ".notified_batches"
LOG_FILE = JOBS_DIR / ".autopipeline.log"
VENV_PYTHON = str(Path.home() / ".openclaw/tools/lda/.venv/bin/python")
LDA_CLI = str(Path.home() / ".openclaw/tools/lda/lda_cli.py")
LEGAL_REF = str(Path.home() / ".openclaw/tools/legal-reference.docx")
MEMO_TEMPLATES = str(Path.home() / ".openclaw/tools/memo-templates.md")
SIGNAL_RPC = "http://127.0.0.1:8080/api/v1/rpc"
USER_NUMBER = "+12063498761"
ENV = {**os.environ, "PATH": "/opt/homebrew/bin:/opt/homebrew/sbin:" + os.environ.get("PATH", "")}

PROGRESS_MIN_INTERVAL = 60

# Legal memo prompt template for CC
# The template reads the memo-templates.md file for structure guidance
MEMO_TEMPLATE_FILE = str(Path.home() / ".openclaw/skills/legal-memo/references/template.md")
DRAFTING_GUIDE_FILE = str(Path.home() / ".openclaw/skills/legal-memo/drafting-guide.md")
DD_SKILL_FILE = str(Path.home() / ".openclaw/skills/legal-due-diligence/SKILL.md")

def build_memo_prompt(instructions, file_list, section_instructions, date):
    """Build the CC prompt for memo generation, reading skill files at runtime."""
    # Read the drafting guide (editable training instructions)
    drafting_guidance = ""
    guide_path = Path(DRAFTING_GUIDE_FILE)
    if guide_path.exists():
        drafting_guidance = (
            "## Drafting Guide\n\n"
            "Follow these drafting instructions carefully:\n\n"
            + guide_path.read_text() + "\n\n"
        )

    # Read the DD review skill for analysis standards
    dd_guidance = ""
    dd_path = Path(DD_SKILL_FILE)
    if dd_path.exists():
        dd_guidance = (
            "## Due Diligence Review Standard\n\n"
            "Follow these analysis standards for document review:\n\n"
            + dd_path.read_text() + "\n\n"
        )

    # Read the memo template for structure reference
    template_guidance = ""
    template_path = Path(MEMO_TEMPLATE_FILE)
    if template_path.exists():
        template_guidance = (
            "## Memo Template Reference\n\n"
            "Follow the Office Memo structure from this template file: "
            + str(template_path) + "\n"
            "Read the file and follow its heading structure, section ordering, and formatting conventions exactly.\n\n"
        )
    else:
        alt_path = Path(MEMO_TEMPLATES)
        if alt_path.exists():
            template_guidance = (
                "## Memo Template Reference\n\n"
                "Follow the Office Memo structure from: " + str(alt_path) + "\n"
                "Read the file and use the Office Memo (Internal) template.\n\n"
            )

    return (
        "You are drafting a legal memorandum analyzing the documents provided.\n\n"
        + drafting_guidance
        + dd_guidance
        + "## Instructions from the attorney:\n"
        + instructions + "\n\n"
        "## Documents to analyze:\n"
        + file_list + "\n\n"
        + template_guidance
        + "## Memo Header\n\n"
        "Use this exact header:\n\n"
        "MEMORANDUM\n\n"
        "TO:\n"
        "[Leave as placeholder -- will be filled by attorney]\n\n"
        "FROM:\n"
        "Haotian Yi, Attorney -- Haotian Legal\n\n"
        "DATE:\n"
        + date + "\n\n"
        "RE:\n"
        "[Descriptive subject line based on the analysis]\n\n"
        "## Section Structure\n\n"
        "I. Executive Summary\n\n"
        "[1-2 paragraph overview of key findings across ALL document categories.]\n\n"
        "**Bottom line:** [One-sentence actionable recommendation.]\n\n"
        + section_instructions + "\n\n"
        "[Final section]. Recommendations / Next Steps\n\n"
        "[Consolidated action items, prioritized.]\n\n"
        "Please do not hesitate to reach out if you have any questions regarding the foregoing.\n\n"
        "Very truly yours,\n\n"
        "___________________________\n"
        "Haotian Yi\n"
        "Attorney\n"
        "Haotian Legal\n\n"
        "## CRITICAL RULES:\n"
        "- Preserve ALL {PLACEHOLDER} tokens exactly as they appear\n"
        "- Do NOT invent or guess real names\n"
        "- Use Roman numeral section headings (I., II., III.)\n"
        "- Use markdown ## for section headings so pandoc can style them\n"
        "- Your response must start IMMEDIATELY with \"MEMORANDUM\" -- no preamble, no commentary\n"
    )


def build_review_comments_prompt(instructions, file_list):
    """Build the CC prompt for contract review with inline comments."""
    # Read the DD review skill for analysis standards
    dd_guidance = ""
    dd_path = Path(DD_SKILL_FILE)
    if dd_path.exists():
        dd_guidance = (
            "## Due Diligence Review Standard\n\n"
            "Follow these analysis standards for document review:\n\n"
            + dd_path.read_text() + "\n\n"
        )

    return (
        "You are a contract review attorney. Review the contract(s) provided and produce "
        "a commented version.\n\n"
        + dd_guidance
        + "## Instructions from the attorney:\n"
        + instructions + "\n\n"
        "## Documents to review:\n"
        + file_list + "\n\n"
        "## Output Format\n\n"
        "For each section or clause of the contract:\n"
        "1. Reproduce the ORIGINAL text verbatim (preserving all formatting and placeholders)\n"
        "2. Immediately below, add your review comments as blockquotes\n\n"
        "Comment format — each comment on its own line:\n"
        "> **[COMMENT]:** explanation of issue/risk/suggestion\n\n"
        "## Review Focus Areas\n\n"
        "Focus your comments on:\n"
        "- **Risks**: Provisions that expose the client to liability or loss\n"
        "- **Ambiguities**: Vague language that could be interpreted unfavorably\n"
        "- **Missing provisions**: Standard protections or clauses that are absent\n"
        "- **Unfavorable terms**: One-sided or below-market provisions\n"
        "- **Suggestions**: Specific language improvements or additions\n\n"
        "## Example Output\n\n"
        "```\n"
        "## 3. Term and Termination\n\n"
        "This Agreement shall commence on the Effective Date and continue for a period "
        "of three (3) years, unless earlier terminated by either party upon thirty (30) "
        "days written notice.\n\n"
        "> **[COMMENT]:** The 30-day notice period is relatively short. Consider "
        "extending to 60 or 90 days to allow adequate transition time.\n\n"
        "> **[COMMENT]:** No termination-for-cause provision. Recommend adding a "
        "clause allowing immediate termination for material breach.\n"
        "```\n\n"
        "## CRITICAL RULES:\n"
        "- Preserve ALL {PLACEHOLDER} tokens exactly as they appear\n"
        "- Do NOT invent or guess real names\n"
        "- Reproduce the original contract text VERBATIM — do not paraphrase or summarize\n"
        "- Use markdown ## for section headings so pandoc can style them\n"
        "- Place comments AFTER each section, not inline within the text\n"
        "- Your response must start IMMEDIATELY with the contract review — no preamble\n"
    )


def build_review_redline_prompt(instructions, file_list):
    """Build the CC prompt for contract review with tracked changes (redline)."""
    # Read the DD review skill for analysis standards
    dd_guidance = ""
    dd_path = Path(DD_SKILL_FILE)
    if dd_path.exists():
        dd_guidance = (
            "## Due Diligence Review Standard\n\n"
            "Follow these analysis standards for document review:\n\n"
            + dd_path.read_text() + "\n\n"
        )

    return (
        "You are a contract review attorney. Review and revise the contract(s) provided, "
        "producing a redlined version that shows all changes.\n\n"
        + dd_guidance
        + "## Instructions from the attorney:\n"
        + instructions + "\n\n"
        "## Documents to review and revise:\n"
        + file_list + "\n\n"
        "## Output Format\n\n"
        "Produce a single document that shows ALL changes using this markup convention:\n"
        "- **Deletions**: Show removed text in ~~strikethrough~~ (e.g., ~~old text~~)\n"
        "- **Additions**: Show new text in **bold** (e.g., **new text**)\n"
        "- Unchanged text remains as-is\n\n"
        "After the redlined document, include a section:\n\n"
        "## Summary of Changes\n\n"
        "List each change made with:\n"
        "- Section/clause reference\n"
        "- Brief description of the change\n"
        "- Rationale for the change\n\n"
        "## Review Focus Areas\n\n"
        "Make specific textual edits to improve the contract. Focus on:\n"
        "- **Ambiguities**: Tighten vague language with precise terms\n"
        "- **Missing protections**: Add standard protective provisions\n"
        "- **Unfavorable terms**: Rebalance one-sided provisions\n"
        "- **Language quality**: Fix inconsistencies, improve clarity\n"
        "- **Legal compliance**: Ensure provisions comply with applicable law\n\n"
        "## Example Output\n\n"
        "```\n"
        "## 3. Term and Termination\n\n"
        "This Agreement shall commence on the Effective Date and continue for a period "
        "of ~~three (3)~~ **five (5)** years, unless earlier terminated by either party "
        "upon ~~thirty (30)~~ **sixty (60)** days written notice. **Notwithstanding the "
        "foregoing, either party may terminate this Agreement immediately upon written "
        "notice if the other party commits a material breach that remains uncured for "
        "thirty (30) days after receipt of written notice thereof.**\n"
        "```\n\n"
        "## CRITICAL RULES:\n"
        "- Preserve ALL {PLACEHOLDER} tokens exactly as they appear\n"
        "- Do NOT invent or guess real names\n"
        "- Use ~~strikethrough~~ for deletions and **bold** for additions CONSISTENTLY\n"
        "- Use markdown ## for section headings so pandoc can style them\n"
        "- Include the Summary of Changes section at the end\n"
        "- Your response must start IMMEDIATELY with the redlined document — no preamble\n"
    )


def build_draft_prompt(instructions, file_list, date):
    """Build the prompt for contract drafting from templates (local Ollama)."""
    return (
        "You are drafting a contract based on the template(s) provided. "
        "Follow the template structure but customize with the specific terms "
        "described in the instructions. Preserve all formatting and section numbering. "
        "Output the complete drafted contract.\n\n"
        "## Instructions from the attorney:\n"
        + instructions + "\n\n"
        "## Template document(s):\n"
        + file_list + "\n\n"
        "## Date:\n"
        + date + "\n\n"
        "## CRITICAL RULES:\n"
        "- Follow the template structure exactly — do not reorganize sections\n"
        "- Replace all template blanks/brackets with the specific terms from the instructions\n"
        "- Preserve ALL {PLACEHOLDER} tokens exactly as they appear (these are anonymized names)\n"
        "- Do NOT invent or guess real names — only use names/terms provided in the instructions\n"
        "- Keep all section numbering, headings, and formatting from the template\n"
        "- Use markdown ## for section headings so pandoc can style them\n"
        "- Your response must start IMMEDIATELY with the drafted contract — no preamble, no commentary\n"
    )


def call_local_llm(prompt, model="qwen3.5-legal"):
    """Call local LLM for privacy-sensitive tasks. Tries MLX first, falls back to Ollama."""
    # Try MLX first (faster, better JSON output)
    try:
        mlx_payload = json.dumps({
            "model": "/Users/claptrap/finetune/mlx-legal",
            "messages": [
                {"role": "system", "content": "/no_think"},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.3,
            "max_tokens": 8192,
        }).encode()
        req = urllib.request.Request(
            "http://127.0.0.1:8801/v1/chat/completions",
            data=mlx_payload,
            headers={"Content-Type": "application/json"}
        )
        resp = urllib.request.urlopen(req, timeout=600)
        data = json.loads(resp.read())
        content = data["choices"][0]["message"]["content"]
        content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
        return content
    except Exception as e:
        logging.warning("MLX failed (%s), falling back to Ollama", e)

    # Fallback to Ollama
    payload = json.dumps({
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "stream": False,
        "options": {"num_ctx": 32768, "num_predict": 8192, "temperature": 0.3}
    }).encode()
    req = urllib.request.Request(
        "http://127.0.0.1:11434/api/chat",
        data=payload,
        headers={"Content-Type": "application/json"}
    )
    resp = urllib.request.urlopen(req, timeout=600)
    data = json.loads(resp.read())
    content = data.get("message", {}).get("content", "")
    content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
    return content


def _get_task_type(meta):
    """Get task type from batch meta, defaulting to 'memo' for backward compatibility."""
    return meta.get("task_type", "memo")


def _task_type_label(task_type):
    """Human-readable label for a task type."""
    labels = {
        "memo": "legal memo",
        "review-comments": "contract review (comments)",
        "review-redline": "contract review (redline)",
        "draft": "contract draft",
    }
    return labels.get(task_type, task_type)


def _build_prompt_for_task(task_type, instructions, file_list, section_instructions=None, date=None):
    """Build the appropriate prompt based on task type."""
    if task_type == "review-comments":
        return build_review_comments_prompt(instructions, file_list)
    elif task_type == "review-redline":
        return build_review_redline_prompt(instructions, file_list)
    elif task_type == "draft":
        return build_draft_prompt(
            instructions, file_list,
            date or datetime.now().strftime("%B %d, %Y"),
        )
    else:
        # Default: memo
        return build_memo_prompt(
            instructions, file_list,
            section_instructions or "",
            date or datetime.now().strftime("%B %d, %Y"),
        )


def _output_filename_prefix(task_type):
    """Return the output filename prefix based on task type."""
    prefixes = {
        "memo": "memo",
        "review-comments": "review_comments",
        "review-redline": "review_redline",
        "draft": "draft",
    }
    return prefixes.get(task_type, "memo")


def log(msg):
    ts = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    with open(LOG_FILE, "a") as f:
        f.write(ts + " " + msg + "\n")


def send_signal(message, attachments=None):
    params = {"recipient": USER_NUMBER, "message": message}
    if attachments:
        params["attachments"] = attachments
    payload = json.dumps({"jsonrpc": "2.0", "method": "send", "params": params, "id": 1}).encode()
    req = urllib.request.Request(SIGNAL_RPC, data=payload, headers={"Content-Type": "application/json"})
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        return json.loads(resp.read())
    except Exception as e:
        log("Signal send failed: " + str(e))
        return None


def progress_bar(done, total, width=20):
    if total == 0:
        return "[" + "\u2591" * width + "] 0%"
    pct = done / total
    filled = int(width * pct)
    return "[" + "\u2588" * filled + "\u2591" * (width - filled) + "] " + str(int(pct * 100)) + "%"


def format_duration(seconds):
    if seconds < 60:
        return str(int(seconds)) + "s"
    elif seconds < 3600:
        m, s = int(seconds // 60), int(seconds % 60)
        return (str(m) + "m " + str(s) + "s") if s else (str(m) + "m")
    else:
        h, m = int(seconds // 3600), int((seconds % 3600) // 60)
        return (str(h) + "h " + str(m) + "m") if m else (str(h) + "h")


def estimate_remaining(elapsed, done, total):
    if done == 0:
        return "calculating..."
    return format_duration((elapsed / done) * (total - done))


def get_completed_batches():
    batches = {}
    for job_dir in JOBS_DIR.iterdir():
        if not job_dir.is_dir() or job_dir.name.startswith(".") or job_dir.name.startswith("unzip"):
            continue
        status_file = job_dir / "status.json"
        if not status_file.exists():
            continue
        try:
            data = json.loads(status_file.read_text())
            bid = data.get("batch_id", "")
            if not bid:
                continue
            if bid not in batches:
                batches[bid] = {"jobs": [], "total": 0, "completed": 0, "failed": 0}
            batches[bid]["total"] += 1
            batches[bid]["jobs"].append(data)
            if data.get("status") == "completed":
                batches[bid]["completed"] += 1
            elif data.get("status") == "failed":
                batches[bid]["failed"] += 1
        except Exception:
            pass
    return {bid: info for bid, info in batches.items()
            if info["completed"] + info["failed"] == info["total"] and info["total"] > 0}


def is_in_file(filepath, value):
    return filepath.exists() and value in filepath.read_text()


def append_to_file(filepath, value):
    with open(filepath, "a") as f:
        f.write(value + "\n")


def get_batch_instructions(batch_id):
    meta_file = META_DIR / (batch_id + ".json")
    if meta_file.exists():
        try:
            return json.loads(meta_file.read_text())
        except Exception:
            pass
    return None


def check_cc_login():
    # Skip login check — let CC fail at runtime with a clear error
    # rather than silently skipping batches. The SSH environment
    # often can't access Keychain, but launchd/terminal context can.
    return True


def extract_folder(input_file_path):
    """Extract the folder path from the job's input_file path.

    For grouped batches, returns the full relative path under unzip dir.
    Example: .../unzip-xxx/TECHNOLOGY & IP/Binar Code/file.pdf -> 'TECHNOLOGY & IP/Binar Code'
             .../unzip-xxx/EXPIRED/file.pdf -> 'EXPIRED'
             .../file.pdf -> '' (root)
    """
    parts = Path(input_file_path).parts
    for i, part in enumerate(parts):
        if part.startswith("unzip-"):
            rel_parts = parts[i + 1:-1]
            if rel_parts:
                return "/".join(rel_parts)
            return ""
    return ""


def extract_top_folder(input_file_path):
    """Extract just the top-level folder name (first level under unzip dir).

    Example: .../unzip-xxx/TECHNOLOGY & IP/subfolder/file.pdf -> 'TECHNOLOGY & IP'
             .../unzip-xxx/file.pdf -> ''
    """
    parts = Path(input_file_path).parts
    for i, part in enumerate(parts):
        if part.startswith("unzip-"):
            rel_parts = parts[i + 1:-1]
            if rel_parts:
                return rel_parts[0]
            return ""
    return ""


def group_jobs_by_folder(jobs):
    """Group completed jobs by their folder within the zip."""
    groups = defaultdict(list)
    for job in jobs:
        if job.get("status") != "completed":
            continue
        folder = extract_folder(job.get("input_file", ""))
        groups[folder].append(job)
    return dict(groups)


def group_jobs_by_top_folder(jobs):
    """Group completed jobs by top-level folder (for cross-batch grouping)."""
    groups = defaultdict(list)
    for job in jobs:
        if job.get("status") != "completed":
            continue
        folder = extract_top_folder(job.get("input_file", ""))
        groups[folder or "(ungrouped)"].append(job)
    return dict(groups)


def merge_mappings(jobs):
    """Merge entity mappings from multiple jobs into one.

    Outputs the format expected by lda_cli.py deanonymize:
    {"mappings": {"{PERSON_1}": {"value": "...", "type": "..."}}, "replacement_log": [...]}
    """
    merged_mappings = {}
    merged_log = []
    for job in jobs:
        mapping_file = JOBS_DIR / job["job_id"] / "mapping.json"
        if not mapping_file.exists():
            continue
        try:
            raw = json.loads(mapping_file.read_text())
            # Collect replacement_log entries
            for entry in raw.get("replacement_log", []):
                merged_log.append(entry)
            # Merge mappings
            entries = raw.get("mappings", raw)
            for placeholder, info in entries.items():
                if placeholder in ("metadata", "replacement_log"):
                    continue
                if isinstance(info, dict):
                    real_value = info.get("value", "")
                    entry_type = info.get("type", "unknown")
                else:
                    real_value = str(info)
                    entry_type = "unknown"
                if placeholder in merged_mappings:
                    existing_val = merged_mappings[placeholder].get("value", "")
                    if real_value and real_value not in existing_val:
                        merged_mappings[placeholder]["value"] = existing_val + " / " + real_value
                else:
                    merged_mappings[placeholder] = {"value": real_value, "type": entry_type}
        except Exception:
            pass
    return {"mappings": merged_mappings, "replacement_log": merged_log}


def remap_placeholders(anon_text, mapping_data, job_suffix):
    """Rename placeholders to be globally unique by adding job suffix.

    {PERSON_1} -> {PERSON_1_J01} to avoid cross-job conflicts.
    """
    remapped_text = anon_text
    remapped_mapping = {}
    remapped_log = []

    mappings = mapping_data.get("mappings", mapping_data)
    for placeholder, info in mappings.items():
        if not placeholder.startswith("{") or placeholder in ("metadata", "replacement_log"):
            continue
        # {PERSON_1} -> {PERSON_1_J01}
        new_placeholder = placeholder[:-1] + "_" + job_suffix + "}"
        remapped_text = remapped_text.replace(placeholder, new_placeholder)
        if isinstance(info, dict):
            remapped_mapping[new_placeholder] = info
        else:
            remapped_mapping[new_placeholder] = {"value": str(info), "type": "unknown"}

    for entry in mapping_data.get("replacement_log", []):
        new_entry = dict(entry)
        old_ph = entry.get("placeholder", "")
        if old_ph.startswith("{") and old_ph.endswith("}"):
            new_entry["placeholder"] = old_ph[:-1] + "_" + job_suffix + "}"
        remapped_log.append(new_entry)

    return remapped_text, {"mappings": remapped_mapping, "replacement_log": remapped_log}



def process_folder_group(folder_name, jobs, instructions, batch_output, task_type="memo"):
    """Process all files in a folder group as a single CC call."""
    safe_folder = folder_name.replace("/", "_").replace(" ", "_") or "root"

    anon_files = []
    for job in jobs:
        anon_file = JOBS_DIR / job["job_id"] / "anonymized.txt"
        if anon_file.exists():
            anon_files.append((job["job_id"], anon_file, job.get("input_filename", "unknown")))

    if not anon_files:
        return {"folder": folder_name, "status": "skipped", "reason": "no anonymized files"}

    file_refs = []
    for i, (job_id, anon_path, orig_name) in enumerate(anon_files, 1):
        file_refs.append("Document " + str(i) + " (job " + job_id + "): " + str(anon_path))

    file_list = "\n".join("- " + ref for ref in file_refs)

    if task_type == "draft":
        # For draft tasks, read template files inline and call local Ollama
        doc_contents = []
        for job_id, anon_path, orig_name in anon_files:
            try:
                content = anon_path.read_text()
                doc_contents.append("### " + orig_name + "\n\n" + content)
            except Exception:
                doc_contents.append("### " + orig_name + "\n\n[Error reading file]")
        file_list_with_content = "\n\n---\n\n".join(doc_contents)
        today = datetime.now().strftime("%B %d, %Y")
        prompt = build_draft_prompt(instructions, file_list_with_content, today)
    elif task_type in ("review-comments", "review-redline"):
        # For review task types, use the specialized prompt
        base_prompt = _build_prompt_for_task(task_type, instructions, file_list)
        prompt = (
            "You are analyzing a group of " + str(len(anon_files)) + " related documents "
            "from the folder '" + folder_name + "'. "
            "These documents are all connected and MUST be analyzed together as a set.\n\n"
            "Read ALL of the following files:\n"
            + file_list + "\n\n"
            + base_prompt
        )
    else:
        # Original memo prompt for folder groups
        prompt = (
            "You are analyzing a group of " + str(len(anon_files)) + " related documents "
            "from the folder '" + folder_name + "'. "
            "These documents are all connected and MUST be analyzed together as a set.\n\n"
            "Read ALL of the following files:\n"
            + file_list
            + "\n\n" + instructions + "\n\n"
            "IMPORTANT:\n"
            "- Analyze all documents in conjunction, not individually.\n"
            "- Identify relationships, cross-references, and dependencies between documents.\n"
            "- Provide conclusions about the group as a whole.\n"
            "- Preserve ALL {PLACEHOLDER} tokens exactly as they appear.\n"
            "- Output ONLY the analysis/memo content, no preamble."
        )

    cc_output = batch_output / (safe_folder + "_cc.md")
    if task_type == "draft":
        # Use local Ollama instead of Claude Code CLI
        try:
            result_text = call_local_llm(prompt)
            if not result_text:
                log("Ollama returned empty response for folder " + folder_name)
                return {"folder": folder_name, "status": "ollama_failed"}
            cc_output.write_text(result_text)
        except Exception as e:
            log("Ollama failed for folder " + folder_name + ": " + str(e))
            return {"folder": folder_name, "status": "ollama_failed"}
    else:
        try:
            cc_result = subprocess.run(
                ["claude", "--print", "--dangerously-skip-permissions", "-p", prompt],
                capture_output=True, text=True, timeout=600, env=ENV,
            )
            if cc_result.returncode != 0:
                log("CC failed for folder " + folder_name + ": " + cc_result.stderr[:200])
                return {"folder": folder_name, "status": "cc_failed"}
            cc_output.write_text(cc_result.stdout)
        except subprocess.TimeoutExpired:
            log("CC timeout for folder " + folder_name)
            return {"folder": folder_name, "status": "cc_timeout"}

    merged = merge_mappings(jobs)
    merged_mapping_file = batch_output / (safe_folder + "_mapping.json")
    merged_mapping_file.write_text(json.dumps(merged, indent=2))

    deanon_file = batch_output / (safe_folder + "_final.md")
    try:
        subprocess.run(
            [VENV_PYTHON, LDA_CLI, "deanonymize",
             "--input", str(cc_output),
             "--mapping", str(merged_mapping_file),
             "--output", str(deanon_file)],
            capture_output=True, text=True, timeout=60,
        )
        if not deanon_file.exists():
            return {"folder": folder_name, "status": "deanon_failed"}
    except Exception:
        return {"folder": folder_name, "status": "deanon_failed"}

    docx_file = batch_output / (safe_folder + ".docx")
    if task_type == "review-redline":
        try:
            from core.tracked_changes import create_redline_docx
            create_redline_docx(deanon_file, docx_file, reference_doc=LEGAL_REF, author="Associate")
            log("Created tracked-changes docx for folder " + folder_name)
        except Exception as e:
            log("Tracked changes failed for " + folder_name + ": " + str(e) + ", falling back to pandoc")
            try:
                subprocess.run(
                    ["pandoc", str(deanon_file), "-o", str(docx_file),
                     "--reference-doc=" + LEGAL_REF],
                    capture_output=True, text=True, timeout=60, env=ENV,
                )
            except Exception:
                pass
    elif task_type == "review-comments":
        try:
            from core.tracked_changes import create_commented_docx
            create_commented_docx(deanon_file, docx_file, reference_doc=LEGAL_REF, author="Associate")
            log("Created commented docx for folder " + folder_name)
        except Exception as e:
            log("Commented docx failed for " + folder_name + ": " + str(e) + ", falling back to pandoc")
            try:
                subprocess.run(
                    ["pandoc", str(deanon_file), "-o", str(docx_file),
                     "--reference-doc=" + LEGAL_REF],
                    capture_output=True, text=True, timeout=60, env=ENV,
                )
            except Exception:
                pass
    else:
        try:
            subprocess.run(
                ["pandoc", str(deanon_file), "-o", str(docx_file),
                 "--reference-doc=" + LEGAL_REF],
                capture_output=True, text=True, timeout=60, env=ENV,
            )
        except Exception:
            pass

    if docx_file.exists() and task_type == "memo":
        format_docx(docx_file)
    output_file = str(docx_file) if docx_file.exists() else str(deanon_file)
    result = {"folder": folder_name, "status": "success", "file": output_file, "doc_count": len(anon_files)}
    # Include clean version for redline tasks
    clean_docx_path = batch_output / (safe_folder + "_clean.docx")
    if clean_docx_path.exists():
        result["clean_file"] = str(clean_docx_path)
    return result



def format_docx(docx_path):
    """Post-process .docx to fix formatting: center MEMORANDUM title, fix spacing."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        # Try LDA venv
        import subprocess, sys
        subprocess.run([VENV_PYTHON, "-c",
            "from docx import Document; "
            "from docx.enum.text import WD_ALIGN_PARAGRAPH; "
            "from docx.shared import Pt; "
            "doc = Document('" + str(docx_path) + "'); "
            "[setattr(p.paragraph_format, 'alignment', WD_ALIGN_PARAGRAPH.CENTER) or "
            " setattr(p.runs[0].font, 'size', Pt(16)) or "
            " setattr(p.runs[0].font, 'bold', True) "
            " for p in doc.paragraphs[:1] if 'MEMORANDUM' in p.text and p.runs]; "
            "doc.save('" + str(docx_path) + "')"],
            capture_output=True, timeout=30)
        return
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        if "MEMORANDUM" in p.text and p.text.strip() == "MEMORANDUM":
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            break
    doc.save(str(docx_path))

def process_unified_memo(all_batch_ids, all_jobs, instructions, output_dir, task_type="memo"):
    """Process multiple batches as a single unified output.

    For memo: Office Memo format with sections per category.
    For review-comments: Contract review with inline blockquote comments.
    For review-redline: Contract review with strikethrough/bold tracked changes.
    For draft: Contract drafting from templates (local Ollama only).
    """
    group_id = all_batch_ids[0]  # Use primary batch as group identifier
    batch_output = output_dir / ("group-" + group_id)
    batch_output.mkdir(parents=True, exist_ok=True)

    completed_jobs = [j for j in all_jobs if j.get("status") == "completed"]
    folder_groups = group_jobs_by_top_folder(completed_jobs)

    # Build file list for the prompt
    file_list_parts = []
    section_instructions_parts = []
    section_num = 2  # Start at II (I is Executive Summary)

    all_remapped_mappings = {"mappings": {}, "replacement_log": []}
    job_counter = 0
    for folder_name in sorted(folder_groups.keys()):
        jobs = folder_groups[folder_name]
        file_list_parts.append("\n### " + folder_name + " (" + str(len(jobs)) + " documents)")
        for job in jobs:
            anon_file = JOBS_DIR / job["job_id"] / "anonymized.txt"
            mapping_file = JOBS_DIR / job["job_id"] / "mapping.json"
            if anon_file.exists() and mapping_file.exists():
                job_counter += 1
                job_suffix = "J" + str(job_counter).zfill(2)
                try:
                    anon_text = anon_file.read_text()
                    raw_mapping = json.loads(mapping_file.read_text())
                    remapped_text, remapped_map = remap_placeholders(anon_text, raw_mapping, job_suffix)
                    remapped_file = batch_output / (job_suffix + "_anon.txt")
                    remapped_file.write_text(remapped_text)
                    if task_type == "draft":
                        # Include file content inline for Ollama (can't read files)
                        orig_name = job.get("input_filename", "document")
                        file_list_parts.append(
                            "\n### " + orig_name + "\n\n" + remapped_text
                        )
                    else:
                        file_list_parts.append("- " + str(remapped_file) + " (job " + job["job_id"] + ")")
                    all_remapped_mappings["mappings"].update(remapped_map["mappings"])
                    all_remapped_mappings["replacement_log"].extend(remapped_map["replacement_log"])
                except Exception as e:
                    file_list_parts.append("- " + str(anon_file) + " (job " + job["job_id"] + ")")

        # Only build section instructions for memo task type
        if task_type == "memo":
            roman = _to_roman(section_num)
            section_instructions_parts.append(
                roman + ". " + folder_name + "\n\n"
                "[Analyze ALL documents in the " + folder_name + " category together. "
                "Identify key issues, obligations, risks, and notable provisions. "
                "Use sub-headings for each major issue found.]\n"
            )
            section_num += 1

    file_list = "\n".join(file_list_parts)
    section_instructions = "\n".join(section_instructions_parts)
    today = datetime.now().strftime("%B %d, %Y")

    # Build prompt based on task type
    prompt = _build_prompt_for_task(task_type, instructions, file_list, section_instructions, today)
    prompt += "\n\nRead ALL files listed above before writing."

    total_files = len(completed_jobs)
    label = _task_type_label(task_type)
    send_signal(
        "Unified " + label + " pipeline started\n"
        "Batches: " + ", ".join(all_batch_ids) + "\n"
        "Files: " + str(total_files) + " across " + str(len(folder_groups)) + " categories\n"
        + "\n".join("  " + k + ": " + str(len(v)) + " docs" for k, v in sorted(folder_groups.items()))
        + "\n\nGenerating " + label + " with " + ("local Ollama..." if task_type == "draft" else "CC...")
    )

    # Single generation call
    prefix = _output_filename_prefix(task_type)
    cc_output = batch_output / (prefix + "_cc.md")
    if task_type == "draft":
        # Use local Ollama for privacy-sensitive drafting
        try:
            result_text = call_local_llm(prompt)
            if not result_text:
                log("Ollama returned empty response for unified draft")
                send_signal("Ollama returned empty response for unified contract draft.")
                return None
            cc_output.write_text(result_text)
        except Exception as e:
            log("Ollama failed for unified draft: " + str(e))
            send_signal("Ollama failed for unified contract draft. Error: " + str(e)[:200])
            return None
    else:
        try:
            cc_result = subprocess.run(
                ["claude", "--print", "--dangerously-skip-permissions", "-p", prompt],
                capture_output=True, text=True, timeout=900, env=ENV,  # 15 min for large unified output
            )
            if cc_result.returncode != 0:
                log("CC failed for unified " + label + ": " + cc_result.stderr[:200])
                send_signal("CC failed for unified " + label + ". Error: " + cc_result.stderr[:200])
                return None
            cc_text = cc_result.stdout
            # Strip any AI preamble before "MEMORANDUM" (memo only)
            if task_type == "memo":
                memo_start = cc_text.find("MEMORANDUM")
                if memo_start > 0:
                    cc_text = cc_text[memo_start:]
            cc_output.write_text(cc_text)
        except subprocess.TimeoutExpired:
            log("CC timeout for unified " + label)
            send_signal("CC timed out generating the unified " + label + " (15 min limit).")
            return None

    engine = "Ollama" if task_type == "draft" else "CC"
    send_signal(engine + " generation complete. Deanonymizing and formatting...")

    # Use pre-built remapped mappings (no cross-job placeholder conflicts)
    merged_mapping_file = batch_output / (prefix + "_mapping.json")
    merged_mapping_file.write_text(json.dumps(all_remapped_mappings, indent=2))

    # Deanonymize
    deanon_file = batch_output / (prefix + "_final.md")
    try:
        subprocess.run(
            [VENV_PYTHON, LDA_CLI, "deanonymize",
             "--input", str(cc_output),
             "--mapping", str(merged_mapping_file),
             "--output", str(deanon_file)],
            capture_output=True, text=True, timeout=60,
        )
        if not deanon_file.exists():
            send_signal("Deanonymization failed.")
            return None
    except Exception as e:
        log("Deanon failed: " + str(e))
        send_signal("Deanonymization failed: " + str(e))
        return None

    # Fallback: replace any remaining placeholders the deanonymizer missed
    try:
        text = deanon_file.read_text()
        remaining = 0
        for placeholder, info in all_remapped_mappings["mappings"].items():
            value = info.get("value", "") if isinstance(info, dict) else str(info)
            if placeholder in text and value:
                text = text.replace(placeholder, value)
                remaining += 1
        if remaining > 0:
            deanon_file.write_text(text)
            log("Fallback replaced " + str(remaining) + " remaining placeholders")
    except Exception:
        pass

    # Format to .docx
    docx_file = batch_output / (prefix + ".docx")
    if task_type == "review-redline":
        try:
            from core.tracked_changes import create_redline_docx
            create_redline_docx(deanon_file, docx_file, reference_doc=LEGAL_REF, author="Associate")
            log("Created tracked-changes docx for unified output")
        except Exception as e:
            log("Tracked changes failed for unified: " + str(e) + ", falling back to pandoc")
            try:
                subprocess.run(
                    ["pandoc", str(deanon_file), "-o", str(docx_file),
                     "--reference-doc=" + LEGAL_REF],
                    capture_output=True, text=True, timeout=60, env=ENV,
                )
            except Exception:
                pass
    elif task_type == "review-comments":
        try:
            from core.tracked_changes import create_commented_docx
            create_commented_docx(deanon_file, docx_file, reference_doc=LEGAL_REF, author="Associate")
            log("Created commented docx for unified output")
        except Exception as e:
            log("Commented docx failed for unified: " + str(e) + ", falling back to pandoc")
            try:
                subprocess.run(
                    ["pandoc", str(deanon_file), "-o", str(docx_file),
                     "--reference-doc=" + LEGAL_REF],
                    capture_output=True, text=True, timeout=60, env=ENV,
                )
            except Exception:
                pass
    else:
        try:
            subprocess.run(
                ["pandoc", str(deanon_file), "-o", str(docx_file),
                 "--reference-doc=" + LEGAL_REF],
                capture_output=True, text=True, timeout=60, env=ENV,
            )
        except Exception:
            pass

    if docx_file.exists() and task_type == "memo":
        format_docx(docx_file)
    output_file = str(docx_file) if docx_file.exists() else str(deanon_file)
    clean_docx_path = batch_output / (prefix + "_clean.docx")
    result = {"status": "success", "file": output_file, "total_files": total_files,
            "categories": len(folder_groups), "task_type": task_type}
    if clean_docx_path.exists():
        result["clean_file"] = str(clean_docx_path)
    return result


def _to_roman(num):
    """Convert integer to Roman numeral."""
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    result = ""
    for val, numeral in vals:
        while num >= val:
            result += numeral
            num -= val
    return result


def run_pipeline(batch_id, batch_info, instructions, task_type="memo"):
    """Run CC pipeline grouped by folder with progress updates (single-batch mode)."""
    batch_output = OUTPUT_DIR / batch_id
    batch_output.mkdir(parents=True, exist_ok=True)

    completed_jobs = [j for j in batch_info["jobs"] if j.get("status") == "completed"]
    folder_groups = group_jobs_by_folder(completed_jobs)
    total_groups = len(folder_groups)
    total_files = len(completed_jobs)
    results = []
    start_time = time.time()

    label = _task_type_label(task_type)
    folder_desc = "\n".join(
        "  " + (name or "(root)") + ": " + str(len(jobs)) + " docs"
        for name, jobs in folder_groups.items()
    )
    send_signal(
        "Auto-pipeline started (" + label + ")\n"
        "Batch: " + batch_id + "\n"
        "Files: " + str(total_files) + " in " + str(total_groups) + " folder(s)\n"
        + folder_desc + "\n\n"
        "Documents in each folder will be " + ("drafted" if task_type == "draft" else "analyzed") + " together.\n"
        + progress_bar(0, total_groups)
    )
    log("Starting pipeline (" + task_type + ") for " + batch_id + " (" + str(total_files) + " files, " + str(total_groups) + " folders)")

    engine = "local Ollama" if task_type == "draft" else "CC"
    action = "Drafting" if task_type == "draft" else "Analyzing"
    for i, (folder_name, jobs) in enumerate(folder_groups.items()):
        display_name = folder_name or "(root)"
        send_signal(
            progress_bar(i, total_groups) + "\n"
            "Processing folder " + str(i + 1) + "/" + str(total_groups) + ": " + display_name + "\n"
            + action + " " + str(len(jobs)) + " documents with " + engine + "..."
        )

        result = process_folder_group(folder_name, jobs, instructions, batch_output, task_type=task_type)
        results.append(result)

        elapsed = time.time() - start_time
        done = i + 1
        status = "\u2713" if result["status"] == "success" else "\u2717"
        eta = estimate_remaining(elapsed, done, total_groups)

        send_signal(
            progress_bar(done, total_groups) + "\n"
            + status + " " + display_name + " (" + str(result.get("doc_count", "?")) + " docs)\n"
            "Elapsed: " + format_duration(elapsed) + " | ETA: ~" + eta
        )

    return results


def deliver_results(batch_id, results, start_time):
    success_files = [r["file"] for r in results if r.get("status") == "success" and r.get("file")]
    fail_count = len([r for r in results if r.get("status") != "success"])
    total_time = format_duration(time.time() - start_time)

    bar = progress_bar(len(results), len(results))
    msg = (
        bar + "\n"
        "Batch complete: " + str(len(success_files)) + "/" + str(len(results)) + " folders processed\n"
        "Total time: " + total_time
    )
    if fail_count:
        msg += "\n\u2717 " + str(fail_count) + " failed"

    if not success_files:
        send_signal(msg)
        return

    # Collect clean files for redline tasks
    clean_files = [r.get("clean_file") for r in results if r.get("clean_file")]
    all_files = success_files + clean_files

    for i in range(0, len(all_files), 5):
        chunk = all_files[i:i + 5]
        chunk_msg = msg if i == 0 else "Files " + str(i + 1) + "-" + str(i + len(chunk)) + " of " + str(len(all_files)) + ":"
        resp = send_signal(chunk_msg, attachments=chunk)
        if resp and resp.get("error"):
            send_signal(chunk_msg + "\nFiles saved at:\n" + "\n".join(chunk))
        time.sleep(1)


def deliver_unified_result(result, batch_ids, start_time):
    """Deliver a unified result (memo, review-comments, or review-redline)."""
    total_time = format_duration(time.time() - start_time)
    if not result or result.get("status") != "success":
        send_signal("Unified pipeline failed after " + total_time)
        return

    label = _task_type_label(result.get("task_type", "memo"))
    msg = (
        progress_bar(1, 1) + "\n"
        "Unified " + label + " complete\n"
        "Batches: " + str(len(batch_ids)) + " | "
        "Files: " + str(result.get("total_files", "?")) + " | "
        "Categories: " + str(result.get("categories", "?")) + "\n"
        "Total time: " + total_time
    )
    attachments = [result["file"]]
    if result.get("clean_file"):
        attachments.append(result["clean_file"])
        msg += "\nIncludes: redlined + clean version"
    resp = send_signal(msg, attachments=attachments)
    if resp and resp.get("error"):
        send_signal(msg + "\nFiles saved at:\n" + "\n".join(attachments))


def resolve_batch_groups(batches, all_completed):
    """Resolve batch groups from meta instructions.

    Returns list of:
      - ("single", batch_id, meta) for standalone batches
      - ("group", [batch_ids], meta) for grouped batches (only once per group)
    """
    seen_groups = set()
    result = []

    for batch_id in batches:
        if is_in_file(PROCESSED_FILE, batch_id):
            continue

        meta = get_batch_instructions(batch_id)
        if not meta or not meta.get("instructions"):
            continue

        # Check if this batch points to a primary batch (it's a related batch)
        primary = meta.get("primary_batch")
        if primary:
            # Skip — will be handled when we encounter the primary batch
            continue

        related = meta.get("related_batches", [])
        if related:
            # This is a primary batch with related batches
            group_key = tuple(sorted([batch_id] + related))
            if group_key in seen_groups:
                continue
            seen_groups.add(group_key)

            # Check all related batches are complete
            all_ready = all(bid in all_completed for bid in [batch_id] + related)
            if not all_ready:
                continue  # Wait for all to complete

            result.append(("group", [batch_id] + related, meta))
        else:
            result.append(("single", batch_id, meta))

    return result


def main():
    META_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    lock_fd = open(LOCK_FILE, "w")
    try:
        fcntl.flock(lock_fd, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except IOError:
        return

    try:
        all_completed = get_completed_batches()
        items = resolve_batch_groups(all_completed, all_completed)

        for item in items:
            if item[0] == "group":
                _, batch_ids, meta = item
                task_type = _get_task_type(meta)
                # Check if primary is already processed
                if is_in_file(PROCESSED_FILE, batch_ids[0]):
                    continue

                # Draft uses local Ollama — no CC login needed
                if task_type != "draft" and not check_cc_login():
                    if not is_in_file(NOTIFIED_FILE, batch_ids[0]):
                        total = sum(all_completed[bid]["completed"] for bid in batch_ids if bid in all_completed)
                        send_signal(
                            "LDA batch group ready (" + str(total) + " files across "
                            + str(len(batch_ids)) + " batches) "
                            "but Claude Code CLI is not logged in.\n"
                            "Please run `claude /login` on the Mac Mini."
                        )
                        append_to_file(NOTIFIED_FILE, batch_ids[0])
                    continue

                start_time = time.time()
                all_jobs = []
                for bid in batch_ids:
                    if bid in all_completed:
                        all_jobs.extend(all_completed[bid]["jobs"])

                total = sum(all_completed.get(bid, {}).get("completed", 0) for bid in batch_ids)
                label = _task_type_label(task_type)
                log("Starting unified " + label + " for " + str(len(batch_ids)) + " batches (" + str(total) + " files)")

                result = process_unified_memo(batch_ids, all_jobs, meta["instructions"], OUTPUT_DIR, task_type=task_type)
                deliver_unified_result(result, batch_ids, start_time)

                # Mark all batches as processed
                for bid in batch_ids:
                    append_to_file(PROCESSED_FILE, bid)

                status = "OK" if result and result.get("status") == "success" else "FAILED"
                log("Unified " + label + " done for " + ",".join(batch_ids) + ": " + status)

            elif item[0] == "single":
                _, batch_id, meta = item
                task_type = _get_task_type(meta)

                # Draft uses local Ollama — no CC login needed
                if task_type != "draft" and not check_cc_login():
                    if not is_in_file(NOTIFIED_FILE, batch_id):
                        info = all_completed[batch_id]
                        send_signal(
                            "LDA batch " + batch_id + " ready (" + str(info["completed"]) + " files) "
                            "but Claude Code CLI is not logged in.\n"
                            "Please run `claude /login` on the Mac Mini."
                        )
                        append_to_file(NOTIFIED_FILE, batch_id)
                    continue

                info = all_completed[batch_id]
                start_time = time.time()
                results = run_pipeline(batch_id, info, meta["instructions"], task_type=task_type)
                deliver_results(batch_id, results, start_time)
                append_to_file(PROCESSED_FILE, batch_id)
                ok = len([r for r in results if r["status"] == "success"])
                log("Pipeline (" + task_type + ") done for " + batch_id + ": " + str(ok) + "/" + str(len(results)) + " folder groups OK")

        # Handle batches with no instructions (notification only)
        for batch_id, info in all_completed.items():
            if is_in_file(PROCESSED_FILE, batch_id) or is_in_file(NOTIFIED_FILE, batch_id):
                continue
            meta = get_batch_instructions(batch_id)
            if meta and (meta.get("instructions") or meta.get("primary_batch")):
                continue  # Has instructions or is part of a group — skip notification
            msg = (
                "LDA batch complete\n"
                "Files: " + str(info["completed"]) + "/" + str(info["total"]) + " anonymized"
            )
            if info["failed"]:
                msg += " (" + str(info["failed"]) + " failed)"
            msg += "\nReply with instructions to start processing."
            send_signal(msg)
            append_to_file(NOTIFIED_FILE, batch_id)
            log("Notified (no instructions) batch " + batch_id)

    except Exception as e:
        log("Autopipeline error: " + str(e))
    finally:
        fcntl.flock(lock_fd, fcntl.LOCK_UN)
        lock_fd.close()


if __name__ == "__main__":
    main()
