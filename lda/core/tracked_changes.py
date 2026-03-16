#!/usr/bin/env python3
"""
Tracked Changes Module

Converts redlined markdown (~~strikethrough~~ deletions and **bold** additions)
into a proper .docx file with real Word Track Changes (revision marks).

Word uses these OOXML elements for tracked changes:
- <w:ins> wraps <w:r> run elements for inserted text
- <w:del> wraps <w:r> with <w:delText> for deleted text
- Both require w:author and w:date attributes
- Regular (unchanged) text uses normal <w:r> elements
"""

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Global revision ID counter
_rev_id = 0


def _next_id():
    """Generate a unique revision ID."""
    global _rev_id
    _rev_id += 1
    return str(_rev_id)


def _reset_ids():
    """Reset the revision ID counter (for testing or new documents)."""
    global _rev_id
    _rev_id = 0


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

# Regex pattern to split markdown text into segments:
#   ~~deleted~~ | **inserted** | plain text
_SEGMENT_RE = re.compile(r'(~~.*?~~|\*\*.*?\*\*)', re.DOTALL)

# Heading pattern: lines starting with # (up to 6 levels)
_HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')

# Summary of Changes sentinel — everything from this heading onward is plain text
_SUMMARY_HEADING_RE = re.compile(r'^#{1,6}\s+Summary\s+of\s+Changes', re.IGNORECASE)


def _parse_segments(text):
    """Parse text into a list of (type, content) tuples.

    Types: 'plain', 'del', 'ins'
    """
    segments = []
    parts = _SEGMENT_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('~~') and part.endswith('~~'):
            inner = part[2:-2]
            if inner:
                segments.append(('del', inner))
        elif part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            if inner:
                segments.append(('ins', inner))
        else:
            if part:
                segments.append(('plain', part))
    return segments


def _split_into_blocks(markdown_text):
    """Split markdown into paragraph blocks, preserving blank-line boundaries.

    Returns a list of dicts:
        {'type': 'heading', 'level': int, 'segments': [...]}
        {'type': 'paragraph', 'segments': [...]}
        {'type': 'summary_heading', 'level': int, 'text': str}
        {'type': 'summary_paragraph', 'text': str}
    """
    blocks = []
    in_summary = False

    # Split on double newlines (paragraph boundaries)
    raw_blocks = re.split(r'\n{2,}', markdown_text.strip())

    for raw in raw_blocks:
        raw = raw.strip()
        if not raw:
            continue

        # Check if this is the Summary of Changes heading
        first_line = raw.split('\n')[0]
        if _SUMMARY_HEADING_RE.match(first_line):
            in_summary = True
            m = _HEADING_RE.match(first_line)
            level = len(m.group(1)) if m else 2
            blocks.append({
                'type': 'summary_heading',
                'level': level,
                'text': m.group(2) if m else first_line,
            })
            # If there's content after the heading in the same block
            rest = '\n'.join(raw.split('\n')[1:]).strip()
            if rest:
                blocks.append({'type': 'summary_paragraph', 'text': rest})
            continue

        if in_summary:
            blocks.append({'type': 'summary_paragraph', 'text': raw})
            continue

        # Check for heading
        m = _HEADING_RE.match(first_line)
        if m and '\n' not in raw.strip():
            # Single-line heading
            level = len(m.group(1))
            heading_text = m.group(2)
            segments = _parse_segments(heading_text)
            blocks.append({'type': 'heading', 'level': level, 'segments': segments})
        else:
            # Regular paragraph (may contain inline line breaks)
            # Collapse single newlines to spaces (markdown convention)
            collapsed = re.sub(r'(?<!\n)\n(?!\n)', ' ', raw)
            segments = _parse_segments(collapsed)
            blocks.append({'type': 'paragraph', 'segments': segments})

    return blocks


# ---------------------------------------------------------------------------
# OOXML tracked change elements
# ---------------------------------------------------------------------------

def _make_run_properties(bold=False, strikethrough=False, underline=False):
    """Create a w:rPr element with optional formatting."""
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if strikethrough:
        strike = OxmlElement('w:strike')
        rPr.append(strike)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    return rPr


def _add_plain_run(paragraph, text):
    """Add a normal (unchanged) run to a paragraph."""
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    paragraph._element.append(run)


def _add_tracked_deletion(paragraph, text, author, date_str):
    """Add a tracked deletion (<w:del>) to a paragraph."""
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), _next_id())
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date_str)

    run = OxmlElement('w:r')
    rPr = _make_run_properties(strikethrough=True)
    run.append(rPr)

    del_text = OxmlElement('w:delText')
    del_text.set(qn('xml:space'), 'preserve')
    del_text.text = text
    run.append(del_text)

    del_elem.append(run)
    paragraph._element.append(del_elem)


def _add_tracked_insertion(paragraph, text, author, date_str):
    """Add a tracked insertion (<w:ins>) to a paragraph."""
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), _next_id())
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date_str)

    run = OxmlElement('w:r')
    rPr = _make_run_properties(underline=True)
    run.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)

    ins_elem.append(run)
    paragraph._element.append(ins_elem)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def _heading_style(level):
    """Map markdown heading level to Word style name."""
    return 'Heading ' + str(min(level, 9))


def _build_tracked_paragraph(doc, segments, author, date_str, style=None):
    """Create a paragraph with tracked change markup from parsed segments."""
    para = doc.add_paragraph()
    if style:
        para.style = style

    for seg_type, content in segments:
        if seg_type == 'plain':
            _add_plain_run(para, content)
        elif seg_type == 'del':
            _add_tracked_deletion(para, content, author, date_str)
        elif seg_type == 'ins':
            _add_tracked_insertion(para, content, author, date_str)

    return para


def create_redline_docx(markdown_path, output_docx_path, reference_doc=None, author="Associate"):
    """Convert redlined markdown to .docx with real Word Track Changes.

    Args:
        markdown_path: Path to the markdown file with ~~deletions~~ and **insertions**.
        output_docx_path: Path for the output .docx file.
        reference_doc: Optional path to a .docx reference document for styling.
        author: Author name for tracked changes (shown in Word's review pane).

    Returns:
        Path to the created .docx file.
    """
    _reset_ids()

    markdown_path = Path(markdown_path)
    output_docx_path = Path(output_docx_path)
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)

    md_text = markdown_path.read_text(encoding='utf-8')
    blocks = _split_into_blocks(md_text)

    # Create document (from reference doc if provided, for styling)
    if reference_doc and Path(reference_doc).exists():
        doc = Document(reference_doc)
        # Remove any existing content from the reference doc
        for para in doc.paragraphs[:]:
            para._element.getparent().remove(para._element)
    else:
        doc = Document()

    date_str = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = block['level']
            try:
                style = doc.styles[_heading_style(level)]
            except KeyError:
                style = None
            _build_tracked_paragraph(doc, block['segments'], author, date_str, style=_heading_style(level))

        elif btype == 'paragraph':
            _build_tracked_paragraph(doc, block['segments'], author, date_str)

        elif btype == 'summary_heading':
            level = block['level']
            doc.add_heading(block['text'], level=level)

        elif btype == 'summary_paragraph':
            doc.add_paragraph(block['text'])

    doc.save(str(output_docx_path))
    return output_docx_path


# ---------------------------------------------------------------------------
# CLI entry point (for testing)
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    import sys
    if len(sys.argv) < 3:
        print("Usage: tracked_changes.py <input.md> <output.docx> [reference.docx] [author]")
        sys.exit(1)

    md_path = sys.argv[1]
    out_path = sys.argv[2]
    ref_doc = sys.argv[3] if len(sys.argv) > 3 else None
    auth = sys.argv[4] if len(sys.argv) > 4 else "Associate"

    result = create_redline_docx(md_path, out_path, reference_doc=ref_doc, author=auth)
    print("Created: " + str(result))


# ---------------------------------------------------------------------------
# Review Comments: Word Comment Annotations
# ---------------------------------------------------------------------------
# Converts markdown with > **[COMMENT]:** blockquotes into a .docx with
# real Word comment boxes (visible in Word's Review tab).

_comment_id_counter = 0


def _next_comment_id():
    """Generate a unique comment ID."""
    global _comment_id_counter
    _comment_id_counter += 1
    return _comment_id_counter


def _reset_comment_ids():
    """Reset the comment ID counter."""
    global _comment_id_counter
    _comment_id_counter = 0


# Pattern to detect comment blockquotes
_COMMENT_RE = re.compile(r'^>\s*\*\*\[COMMENT\]:\*\*\s*(.+)$', re.MULTILINE)


def _get_or_create_comments_part(document):
    """Get or create the comments XML part in the document package."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    COMMENTS_RELTYPE = (
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    )

    # Check if comments part already exists
    for rel in document.part.rels.values():
        if "comments" in rel.reltype:
            return rel.target_part

    # Create a new comments part
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"'
        ' xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
        ' xmlns:mv="urn:schemas-microsoft-com:mac:vml"'
        ' xmlns:o="urn:schemas-microsoft-com:office:office"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        ' xmlns:v="urn:schemas-microsoft-com:vml"'
        ' xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:w10="urn:schemas-microsoft-com:office:word"'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"'
        ' xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"'
        ' xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"'
        ' xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' mc:Ignorable="w14 wp14">'
        '</w:comments>'
    )
    comments_part = Part(
        PackURI('/word/comments.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
        comments_xml.encode('utf-8'),
        document.part.package,
    )
    document.part.relate_to(comments_part, COMMENTS_RELTYPE)
    return comments_part


def _add_comment_to_paragraph(document, paragraph, comment_text, author, date_str,
                              comments_part):
    """Attach a Word comment annotation to a paragraph.

    Creates:
    - A <w:comment> element in the comments part
    - <w:commentRangeStart> / <w:commentRangeEnd> around the paragraph content
    - <w:commentReference> run after the range end
    """
    from lxml import etree

    comment_id = str(_next_comment_id())

    # --- 1. Add <w:comment> to the comments part XML ---
    comments_elem = etree.fromstring(comments_part.blob)
    comment_elem = etree.SubElement(comments_elem, qn('w:comment'))
    comment_elem.set(qn('w:id'), comment_id)
    comment_elem.set(qn('w:author'), author)
    comment_elem.set(qn('w:date'), date_str)
    comment_elem.set(qn('w:initials'), author[0] if author else 'A')

    # Comment text as a paragraph inside the comment element
    cp = etree.SubElement(comment_elem, qn('w:p'))
    cr = etree.SubElement(cp, qn('w:r'))
    ct = etree.SubElement(cr, qn('w:t'))
    ct.set(qn('xml:space'), 'preserve')
    ct.text = comment_text

    # Persist updated XML back to the part
    comments_part._blob = etree.tostring(
        comments_elem, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # --- 2. Wrap paragraph content with comment range markers ---
    p_elem = paragraph._element

    # commentRangeStart — insert before all existing children
    range_start = OxmlElement('w:commentRangeStart')
    range_start.set(qn('w:id'), comment_id)
    p_elem.insert(0, range_start)

    # commentRangeEnd — append after existing content
    range_end = OxmlElement('w:commentRangeEnd')
    range_end.set(qn('w:id'), comment_id)
    p_elem.append(range_end)

    # commentReference run — must follow range end
    ref_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'CommentReference')
    rPr.append(rStyle)
    ref_run.append(rPr)
    ref_mark = OxmlElement('w:commentReference')
    ref_mark.set(qn('w:id'), comment_id)
    ref_run.append(ref_mark)
    p_elem.append(ref_run)


def _parse_review_markdown(md_text):
    """Parse review-comments markdown into blocks.

    Returns a list of dicts:
        {'type': 'heading', 'level': int, 'text': str}
        {'type': 'paragraph', 'text': str}
        {'type': 'comment', 'text': str}
    
    Comment blocks are > **[COMMENT]:** lines. They attach to the preceding
    non-comment block.
    """
    blocks = []
    lines = md_text.strip().split('\n')
    current_para_lines = []

    def flush_para():
        if current_para_lines:
            text = ' '.join(current_para_lines).strip()
            if text:
                blocks.append({'type': 'paragraph', 'text': text})
            current_para_lines.clear()

    for line in lines:
        stripped = line.strip()

        # Check for comment blockquote
        cm = _COMMENT_RE.match(stripped)
        if cm:
            flush_para()
            blocks.append({'type': 'comment', 'text': cm.group(1).strip()})
            continue

        # Skip blank lines (paragraph separator)
        if not stripped:
            flush_para()
            continue

        # Check for heading
        hm = _HEADING_RE.match(stripped)
        if hm:
            flush_para()
            level = len(hm.group(1))
            blocks.append({'type': 'heading', 'level': level, 'text': hm.group(2).strip()})
            continue

        # Regular text line — accumulate into current paragraph
        current_para_lines.append(stripped)

    flush_para()
    return blocks


def create_commented_docx(markdown_path, output_docx_path, reference_doc=None,
                          author="Associate"):
    """Convert review markdown with blockquote comments to .docx with real Word comments.

    Parses > **[COMMENT]:** lines as Word comment annotations attached to the
    preceding paragraph.

    Args:
        markdown_path: Path to the markdown file with comment blockquotes.
        output_docx_path: Path for the output .docx file.
        reference_doc: Optional path to a .docx reference document for styling.
        author: Author name for comments (shown in Word's Review pane).

    Returns:
        Path to the created .docx file.
    """
    from lxml import etree

    _reset_comment_ids()

    markdown_path = Path(markdown_path)
    output_docx_path = Path(output_docx_path)
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)

    md_text = markdown_path.read_text(encoding='utf-8')
    blocks = _parse_review_markdown(md_text)

    # Create document
    if reference_doc and Path(reference_doc).exists():
        doc = Document(reference_doc)
        # Remove existing content from the reference doc
        for para in doc.paragraphs[:]:
            para._element.getparent().remove(para._element)
    else:
        doc = Document()

    date_str = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    # Get or create the comments part once
    comments_part = _get_or_create_comments_part(doc)

    # Track the last non-comment paragraph so we can attach comments to it
    last_para = None

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = block['level']
            last_para = doc.add_heading(block['text'], level=min(level, 9))

        elif btype == 'paragraph':
            last_para = doc.add_paragraph(block['text'])

        elif btype == 'comment':
            if last_para is not None:
                _add_comment_to_paragraph(
                    doc, last_para, block['text'], author, date_str, comments_part
                )
            else:
                # No preceding paragraph — add as a standalone note paragraph
                note_para = doc.add_paragraph('[Comment] ' + block['text'])
                last_para = note_para

    doc.save(str(output_docx_path))
    return output_docx_path


def create_clean_revised(markdown_path, output_path):
    """Strip redline markup to produce a clean revised document.
    
    Removes ~~deleted text~~ entirely and unwraps **added text** to plain text.
    The result is the 'accepted all changes' version.
    """
    import re
    text = Path(markdown_path).read_text(encoding='utf-8')
    
    # Remove deletions (~~text~~)
    text = re.sub(r'~~(.*?)~~', '', text, flags=re.DOTALL)
    
    # Unwrap insertions (**text**) — but preserve markdown headings
    # Only strip bold markers that are inline (not at start of line after #)
    lines = text.split('\n')
    result = []
    for line in lines:
        if line.lstrip().startswith('#'):
            result.append(line)
        else:
            # Remove **bold** markers (insertions become plain text)
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            result.append(line)
    
    text = '\n'.join(result)
    
    # Clean up double spaces left by removed deletions
    text = re.sub(r'  +', ' ', text)
    # Clean up empty lines left by removed content
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    Path(output_path).write_text(text, encoding='utf-8')
